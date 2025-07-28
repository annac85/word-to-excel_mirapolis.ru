# =====================================================
# Конвертер Word → Excel для Мираполис
# Автор: Анна Черкасова (https://cherkasovaanna.ru/)
# 
# ⚠️ Использование в коммерческих целях — 
#    только с письменного разрешения автора.
#    Контакты: anna@cherkasovaanna.ru | ТГ @annac1119
# =====================================================


from docx import Document
import pandas as pd
import os
import re


def extract_answers_from_word(doc):
    """
    Собирает таблицу вопросов и правильных ответов из Word в шаблоне платформы Мираполис.
    Обрабатывает несколько строк вопросов и ответов.
    :param doc: Word-документ.
    :return: Словарь, где ключ — номер вопроса, значение — список правильных ответов.
    """
    answers = {}
    for table in doc.tables:
        rows = table.rows
        for i in range(0, len(rows), 2):  # Проходит парами строк: вопрос -> ответ
            if i + 1 < len(rows):  # Проверка, что есть пара строк
                question_row = rows[i].cells
                answer_row = rows[i + 1].cells

                # Сопоставляем вопросы и ответы
                for j in range(max(len(question_row), len(answer_row))):
                    question_text = question_row[j].text.strip() if j < len(question_row) else ""
                    answer_text = answer_row[j].text.strip() if j < len(answer_row) else ""

                    if question_text.isdigit():  # Если текст — номер вопроса
                        question_number = int(question_text)
                        try:
                            correct_answers = list(map(int, answer_text.split(',')))
                            answers[question_number] = correct_answers
                        except ValueError:
                            print(f"Ошибка обработки ответов: {answer_text}")
    return answers


def clean_answer(answer):
    """
    Убирает номера ответов из текста (например, '1)' или '2)').
    :param answer: Текст ответа.
    :return: Очищенный текст.
    """
    return re.sub(r'^\d+\)\s*', '', answer).strip()


def extract_questions_answers(docx_file, output_excel):
    """
    Обрабатывает документ Word: извлекает вопросы, ответы и ключи.
    :param docx_file: Путь к Word файлу.
    :param output_excel: Путь к выходному Excel файлу.
    """
    try:
        # Открываем Word файл
        document = Document(docx_file)
        if not document.tables:
            print(f"Файл {docx_file} не содержит таблиц.")
            return

        # Извлекаем ключи (правильные ответы)
        keys = extract_answers_from_word(document)

        # Имя файла без расширения как название группы
        group_name = os.path.basename(docx_file).replace(".docx", "")

        # Списки для хранения результатов
        rows = []
        current_question = None
        question_number = 0  # Номер текущего вопроса
        answer_number = 0    # Номер текущего ответа для каждого вопроса

        # Считываем содержимое файла
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                # Проверяем, начинается ли текст с числа (вопрос)
                match = re.match(r'^(\d+)\.', text)
                if match:
                    current_question = text
                    question_number += 1
                    answer_number = 0  # Сброс номера ответа для нового вопроса
                else:
                    # Если это ответ, добавляем его с текущим вопросом в отдельную строку
                    if current_question:
                        answer_number += 1
                        correctness = 0  # По умолчанию ответ неверный
                        question_type = 0  # По умолчанию тип одиночного выбора
                        if question_number in keys:  # Проверяем наличие ключа для текущего вопроса
                            if answer_number in keys[question_number]:
                                correctness = 1
                            if len(keys[question_number]) > 1:
                                question_type = 4  # Тип для множественного выбора

                        # Если это первый ответ, добавляем все данные вопроса
                        if answer_number == 1:
                            rows.append({
                                "Название группы": group_name,
                                "Код": "",
                                "Название вопроса": current_question,
                                "Тип": question_type,
                                "Категория": "",
                                "Время": "",
                                "Текст": current_question,
                                "Картинка": "",
                                "Разрешить комментарий к ответу": 1,
                                "Баллы": 1,
                                "Считать по ответам": 0,
                                "Попыток": "",
                                "Сообщение при правильном ответе": "",
                                "Сообщение при неправильном ответе": "",
                                "Порядок ответов": 1,
                                "Кол-во столбцов": 1,
                                "Использовать нейтральные": 0,
                                "Макс. кол-во ответов": "",
                                "Ответ": clean_answer(text),
                                "Картинка ответа": "",
                                "Балл": 1,
                                "Правильность": correctness,
                                "Комментарий": ""
                            })
                        else:
                            # Остальные строки заполняем только ответом
                            rows.append({
                                "Название группы": "",
                                "Код": "",
                                "Название вопроса": "",
                                "Тип": "",
                                "Категория": "",
                                "Время": "",
                                "Текст": "",
                                "Картинка": "",
                                "Разрешить комментарий к ответу": "",
                                "Баллы": "",
                                "Считать по ответам": "",
                                "Попыток": "",
                                "Сообщение при правильном ответе": "",
                                "Сообщение при неправильном ответе": "",
                                "Порядок ответов": "",
                                "Кол-во столбцов": "",
                                "Использовать нейтральные": "",
                                "Макс. кол-во ответов": "",
                                "Ответ": clean_answer(text),
                                "Картинка ответа": "",
                                "Балл": "",
                                "Правильность": correctness,
                                "Комментарий": ""
                            })

        # Создает DataFrame
        data = pd.DataFrame(rows)

        # Сохраняет в Excel
        data.to_excel(output_excel, index=False)
        print(f"Данные успешно сохранены в {output_excel}")
    except Exception as e:
        print(f"Ошибка обработки файла {docx_file}: {e}")


if __name__ == "__main__":
    # Указать папку
    input_folder = r"C:\Users"

    # Проходит по всем Word файлам в указанной папке
    for filename in os.listdir(input_folder):
        # Игнорирует временные файлы Word
        if filename.endswith(".docx") and not filename.startswith("~$"):
            input_path = os.path.join(input_folder, filename)
            output_path = os.path.join(input_folder, filename.replace(".docx", ".xlsx"))
            try:
                extract_questions_answers(input_path, output_path)
            except Exception as e:
                print(f"Общая ошибка при обработке файла {filename}: {e}")
